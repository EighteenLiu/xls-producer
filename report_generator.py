# -*- coding: utf-8 -*-
"""Generate monthly Word reports from the street summary workbook.

The generator uses the existing DOCX samples as templates.  It supports explicit
``{{PLACEHOLDER}}`` tokens, and also works with the original sample wording by
replacing the known date/report paragraphs in place.
"""

from __future__ import annotations

import argparse
import io
import math
import re
import shutil
import tempfile
import time as time_module
import zipfile
from dataclasses import dataclass
from datetime import datetime, time
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import Workbook, load_workbook
from openpyxl.utils.datetime import from_excel


BASE_DIR = Path(__file__).resolve().parent
DEFAULT_TEMPLATE_DIR = BASE_DIR / "data" / "report_templates"

STREETS = [
    "德胜街道",
    "什刹海街道",
    "西长安街街道",
    "大栅栏街道",
    "天桥街道",
    "新街口街道",
    "金融街街道",
    "椿树街道",
    "陶然亭街道",
    "展览路街道",
    "月坛街道",
    "广内街道",
    "牛街街道",
    "白纸坊街道",
    "广外街道",
]

CHART_NS = {
    "c": "http://schemas.openxmlformats.org/drawingml/2006/chart",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
    "c14": "http://schemas.microsoft.com/office/drawing/2007/8/2/chart",
    "c15": "http://schemas.microsoft.com/office/drawing/2012/chart",
    "c16": "http://schemas.microsoft.com/office/drawing/2014/chart",
    "wps": "https://web.wps.cn/et/2018/main",
}
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
for prefix, uri in CHART_NS.items():
    ET.register_namespace(prefix, uri)


def clean_text(value) -> str:
    text = "" if value is None else str(value)
    return re.sub(r"\s+", "", text).strip()


def number(value) -> float:
    if value in (None, ""):
        return 0.0
    if isinstance(value, bool):
        return float(value)
    if isinstance(value, (int, float)):
        if isinstance(value, float) and math.isnan(value):
            return 0.0
        return float(value)
    try:
        return float(str(value).strip())
    except ValueError:
        return 0.0


def to_datetime(value) -> datetime | None:
    if value is None:
        return None
    if isinstance(value, datetime):
        return value
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        try:
            return from_excel(value)
        except Exception:
            return None
    if isinstance(value, str):
        text = value.strip()
        for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d"):
            try:
                parsed = datetime.strptime(text, fmt)
                return datetime.combine(parsed.date(), time.min)
            except ValueError:
                pass
    return None


def parse_date(value: str, end_of_day: bool = False) -> datetime:
    parsed = datetime.strptime(value, "%Y-%m-%d")
    return datetime.combine(parsed.date(), time.max.replace(microsecond=0) if end_of_day else time.min)


def format_period(start_date: datetime, end_date: datetime) -> str:
    return f"{start_date.year}年{start_date.month}月{start_date.day}日至{end_date.year}年{end_date.month}月{end_date.day}日"


def format_month(end_date: datetime) -> str:
    return f"{end_date.year}年{end_date.month}月份"


def fmt_int(value: float | int) -> str:
    return str(int(round(value)))


def fmt_pct(value: float, digits: int = 2) -> str:
    return f"{value * 100:.{digits}f}%"


def with_ratio(items: list[tuple[str, float]], total: float) -> str:
    return "、".join(f"{name}（{fmt_int(value)}个，占比{fmt_pct(value / total if total else 0)}）" for name, value in items)


def list_counts(items: list[tuple[str, float]], unit: str = "处") -> str:
    return "、".join(f"{name}{fmt_int(value)}{unit}" for name, value in items)


def ratio_items(items: list[tuple[str, float]], total: float, unit: str = "处") -> str:
    return "、".join(f"{name}（{fmt_int(value)}{unit}，占比{fmt_pct(value / total if total else 0)}）" for name, value in items)


def ranked(values: dict[str, float], reverse: bool = True, count: int = 3, include_zero: bool = True) -> list[tuple[str, float]]:
    order = {street: idx for idx, street in enumerate(STREETS)}
    rows = [(street, values.get(street, 0.0)) for street in STREETS]
    if not include_zero:
        rows = [(street, value) for street, value in rows if value > 0]
    return sorted(rows, key=lambda item: ((-item[1] if reverse else item[1]), order[item[0]]))[:count]


@dataclass
class SummaryData:
    summary_xlsx: Path
    sheet1_values: dict[tuple[str, str], dict[str, float]]
    sheet2_rows: dict[str, dict[str, float]]
    missing: list[str]

    @classmethod
    def load(cls, summary_xlsx: Path) -> "SummaryData":
        wb = load_workbook(summary_xlsx, data_only=True)
        if "Sheet1" not in wb.sheetnames:
            raise ValueError("汇总表缺少 Sheet1")
        ws = wb["Sheet1"]
        streets = [clean_text(ws.cell(1, col).value) for col in range(4, ws.max_column + 1)]
        current_category = ""
        sheet1_values: dict[tuple[str, str], dict[str, float]] = {}
        for row in range(2, ws.max_row + 1):
            category = clean_text(ws.cell(row, 2).value)
            metric = clean_text(ws.cell(row, 3).value)
            if category:
                current_category = category
            if not metric:
                continue
            values = {}
            for offset, street in enumerate(streets, start=4):
                if street:
                    values[street] = number(ws.cell(row, offset).value)
            sheet1_values[(current_category, metric)] = values

        sheet2_rows: dict[str, dict[str, float]] = {}
        if "Sheet2" in wb.sheetnames:
            ws2 = wb["Sheet2"]
            headers = [clean_text(ws2.cell(1, col).value) for col in range(1, ws2.max_column + 1)]
            for row in range(2, ws2.max_row + 1):
                street = clean_text(ws2.cell(row, 1).value)
                if not street:
                    continue
                sheet2_rows[street] = {
                    headers[col - 1]: number(ws2.cell(row, col).value)
                    for col in range(2, ws2.max_column + 1)
                }
        return cls(summary_xlsx=summary_xlsx, sheet1_values=sheet1_values, sheet2_rows=sheet2_rows, missing=[])

    def values(self, category: str, metric: str, required: bool = True) -> dict[str, float]:
        key = (clean_text(category), clean_text(metric))
        if key not in self.sheet1_values:
            if required:
                self.missing.append(f"Sheet1 缺少指标：{category} / {metric}")
            return {street: 0.0 for street in STREETS}
        return {street: self.sheet1_values[key].get(street, 0.0) for street in STREETS}

    def total(self, category: str, metric: str) -> float:
        return sum(self.values(category, metric).values())

    def top_issue_metrics(self, metrics: Iterable[tuple[str, str]], count: int = 6) -> list[tuple[str, float]]:
        rows = []
        for category, metric in metrics:
            total = self.total(category, metric)
            if total > 0:
                rows.append((metric, total))
        return sorted(rows, key=lambda item: item[1], reverse=True)[:count]

    def category_issue_totals(self, category: str, count: int = 15, include_zero: bool = True) -> list[tuple[str, float]]:
        wanted = clean_text(category)
        rows = []
        for (row_category, metric), values in self.sheet1_values.items():
            if row_category != wanted or metric in ("合计", "社会单位检查问题数", "餐饮单位检查问题数"):
                continue
            total = sum(values.get(street, 0.0) for street in STREETS)
            if include_zero or total > 0:
                rows.append((metric, total))
        return sorted(rows, key=lambda item: item[1], reverse=True)[:count]


def load_ledger_counts(source_ledger_xlsx: Path, start_date: datetime, end_date: datetime) -> dict[str, dict[str, int]]:
    wb = load_workbook(source_ledger_xlsx, data_only=True, read_only=True)
    result: dict[str, dict[str, int]] = {"社会单位": {}, "餐饮单位": {}}
    for sheet_name, name_col in (("社会单位", 3), ("餐饮单位", 3)):
        if sheet_name not in wb.sheetnames:
            continue
        ws = wb[sheet_name]
        seen: dict[str, set[str]] = {street: set() for street in STREETS}
        for row in ws.iter_rows(min_row=2, values_only=True):
            checked_at = to_datetime(row[0] if row else None)
            if checked_at is None or not (start_date <= checked_at <= end_date):
                continue
            street = clean_text(row[1] if len(row) > 1 else "")
            if street not in seen:
                continue
            unit_name = clean_text(row[name_col - 1] if len(row) >= name_col else "")
            seen[street].add(unit_name or f"row-{ws._current_row}")
        result[sheet_name] = {street: len(names) for street, names in seen.items()}
    return result


def load_ledger_issue_totals(source_ledger_xlsx: Path, start_date: datetime, end_date: datetime) -> dict[str, list[tuple[str, float]]]:
    wb = load_workbook(source_ledger_xlsx, data_only=True, read_only=True)
    result: dict[str, list[tuple[str, float]]] = {"社会单位": [], "餐饮单位": []}
    skip_headers = {
        "检查日期",
        "街道名称",
        "社会单位名称",
        "餐饮单位名称",
        "合计",
        "整改情况（已整改：0，未整改：1）",
    }
    for sheet_name in ("社会单位", "餐饮单位"):
        if sheet_name not in wb.sheetnames:
            continue
        ws = wb[sheet_name]
        headers = [clean_text(value) for value in next(ws.iter_rows(min_row=1, max_row=1, values_only=True))]
        issue_columns = [
            (index, header)
            for index, header in enumerate(headers)
            if header and header not in skip_headers and "整改情况" not in header
        ]
        totals = {header: 0.0 for _, header in issue_columns}
        for row in ws.iter_rows(min_row=2, values_only=True):
            checked_at = to_datetime(row[0] if row else None)
            if checked_at is None or not (start_date <= checked_at <= end_date):
                continue
            for index, header in issue_columns:
                totals[header] = totals.get(header, 0.0) + number(row[index] if index < len(row) else None)
        result[sheet_name] = sorted(
            ((header, total) for header, total in totals.items() if total > 0),
            key=lambda item: item[1],
            reverse=True,
        )
    return result


RESIDENTIAL_CHART_METRICS = [
    ("TOTAL_ISSUES_BY_STREET", "各街道问题数"),
    ("MAIN_ISSUE_TYPES", "主要问题类型"),
    ("MAIN_ISSUE_TYPES", "检查发现问题类型"),
    ("RESIDENT_ACCURACY", "各街道居民自主投放准确率"),
    ("PURITY_RATE", "各街道桶内垃圾分类纯净率"),
    ("DIRTY_OVERFLOW", "各街道桶站满冒环境脏乱数"),
    ("PEAK_UNCOVERED", "各街道高峰时段未开盖数"),
    ("FACILITY_BUILD", "各街道桶站建设不规范数"),
    ("TOTAL_ISSUES_BY_STREET", "非值守时段各街道问题数"),
    ("DIRTY_OVERFLOW", "非值守时段桶站满冒环境脏乱数"),
    ("FACILITY_BUILD", "非值守时段桶站建设不规范数"),
]

SOCIAL_CHART_METRICS = [
    ("SOCIAL_ISSUES_BY_STREET", "社会单位问题数"),
    ("SOCIAL_ISSUE_TYPES", "社会单位生活垃圾分类日常检查各类问题数"),
    ("RESTAURANT_ISSUES_BY_STREET", "餐饮单位问题数"),
    ("RESTAURANT_ISSUE_TYPES", "餐饮单位生活垃圾分类日常检查各类问题数"),
]


def build_chart_datasets(
    summary: SummaryData,
    ledger_issue_totals: dict[str, list[tuple[str, float]]] | None = None,
) -> dict[str, list[tuple[str, float]]]:
    ledger_issue_totals = ledger_issue_totals or {}
    construction = summary.values("分类设施建设达标情况", "合计")
    management = summary.values("分类设施管理达标情况", "合计")
    resident_bad = summary.values("居民自主投放情况", "居民自主投放不准确")
    total_by_street = {
        street: construction[street] + management[street] + resident_bad[street]
        for street in STREETS
    }
    social = summary.values("社会单位检查情况", "社会单位检查问题数")
    restaurant = summary.values("餐饮单位检查情况", "餐饮单位检查问题数")
    datasets = {
        "TOTAL_ISSUES_BY_STREET": ranked(total_by_street, reverse=False, count=len(STREETS)),
        "MAIN_ISSUE_TYPES": summary.top_issue_metrics(
            [
                ("居民自主投放情况", "居民自主投放不准确"),
                ("分类设施管理达标情况", "桶站周边不洁"),
                ("分类设施管理达标情况", "桶站地面脏污"),
                ("分类设施管理达标情况", "容器满冒"),
                ("分类设施管理达标情况", "垃圾积存"),
                ("分类设施管理达标情况", "桶内分类不纯净"),
                ("分类设施建设达标情况", "高峰时段未开盖"),
                ("分类设施建设达标情况", "合计"),
            ]
        ),
        "RESIDENT_ACCURACY": ranked(summary.values("居民自主投放情况", "投放准确率"), reverse=False, count=len(STREETS)),
        "PURITY_RATE": ranked(summary.values("纯净率", "桶内分类纯净率"), reverse=False, count=len(STREETS)),
        "DIRTY_OVERFLOW": ranked({
            street: summary.values("分类设施管理达标情况", "桶站周边不洁")[street]
            + summary.values("分类设施管理达标情况", "桶站地面脏污")[street]
            + summary.values("分类设施管理达标情况", "容器满冒")[street]
            + summary.values("分类设施管理达标情况", "垃圾积存")[street]
            for street in STREETS
        }, reverse=False, count=len(STREETS)),
        "PEAK_UNCOVERED": ranked(summary.values("分类设施建设达标情况", "高峰时段未开盖"), reverse=False, count=len(STREETS)),
        "FACILITY_BUILD": ranked(construction, reverse=False, count=len(STREETS)),
        "SOCIAL_ISSUES_BY_STREET": ranked(social, reverse=False, count=len(STREETS)),
        "RESTAURANT_ISSUES_BY_STREET": ranked(restaurant, reverse=False, count=len(STREETS)),
        "SOCIAL_ISSUE_TYPES": (
            ledger_issue_totals.get("社会单位") or summary.category_issue_totals("社会单位检查情况", count=15, include_zero=False)
        )[:15],
        "RESTAURANT_ISSUE_TYPES": (
            ledger_issue_totals.get("餐饮单位") or summary.category_issue_totals("餐饮单位检查情况", count=15, include_zero=False)
        )[:15],
    }
    return datasets


def replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    if not paragraph.runs:
        return
    original = paragraph.text
    replaced = original
    for old, new in replacements.items():
        replaced = replaced.replace(old, new)
    if replaced == original:
        return
    for run in paragraph.runs:
        run.text = ""
    paragraph.runs[0].text = replaced


def replace_text(doc: Document, replacements: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(8)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def fill_table(table, headers: list[str], rows: list[list[str | int | float]]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    while len(table.columns) < len(headers):
        table.add_column(Pt(42))
    while len(table.columns) > len(headers):
        break
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_text(cell, header, bold=True)
        shade_cell(cell, "D9EAF7")
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            set_cell_text(cells[idx], "" if value is None else str(value))


def social_table_rows(summary: SummaryData, counts: dict[str, int]) -> tuple[list[str], list[list[str | int]]]:
    headers = [
        "街道名称",
        "社会单位数量",
        "无培训活动、会议记录、照片等培训材料",
        "无适量点餐、光盘行动等宣传内容",
        "单位无容器配置或容器配置不全",
        "公共场所区域未成组设置可回收物和其他容器数（组数）",
        "容器无便利性措施",
        "容器标识不合格",
        "无宣传氛围",
        "无厨余垃圾收运合同或不合格",
        "无其他垃圾收运合同或不合格",
        "无可回收物收运合同或不合格",
        "无有害垃圾收运合同或不合格",
        "无废弃油脂合同或不合格",
        "无垃圾分类工作方案",
        "无四分类垃圾清运台账或四分类清运台账不合格",
        "无油水分离装置",
        "合计",
    ]
    metrics = [
        "无培训活动、会议记录、照片等培训材料",
        "无适量点餐、光盘行动等宣传内容",
        "单位无容器配置或容器配置不全",
        "公共场所区域(办公楼外区域、办事大厅等)未成组设置可回收物和其他容器数（组数）",
        "容器无便利性措施",
        "容器标识不合格数",
        "无宣传氛围",
        "无厨余垃圾收运合同或不合格",
        "无其他垃圾收运合同或不合格",
        "无可回收物收运合同或不合格",
        "无有害垃圾收运合同或不合格",
        "无废弃油脂合同或不合格",
        "无垃圾分类工作方案",
        "无四分类垃圾清运台账或四分类清运台账不合格",
        "无油水分离装置",
    ]
    metric_values = [summary.values("社会单位检查情况", metric, required=False) for metric in metrics]
    total_values = summary.values("社会单位检查情况", "社会单位检查问题数")
    rows = [["总计", sum(counts.values()), *[fmt_int(sum(values.values())) for values in metric_values], fmt_int(sum(total_values.values()))]]
    for street, _value in ranked(total_values, reverse=True, count=len(STREETS), include_zero=True):
        rows.append([street, counts.get(street, 0), *[fmt_int(values[street]) for values in metric_values], fmt_int(total_values[street])])
    return headers, rows


def restaurant_table_rows(summary: SummaryData, counts: dict[str, int]) -> tuple[list[str], list[list[str | int]]]:
    headers = [
        "街道名称",
        "餐饮单位数量",
        "无适量点餐、光盘行动等宣传内容",
        "集中用餐区未成组设置厨余和其他垃圾容器数",
        "容器无标识或标识不合格",
        "无宣传氛围",
        "无厨余垃圾收运合同或不合格",
        "无其他垃圾收运合同或不合格",
        "无厨余垃圾排放登记方式",
        "无非居民其他垃圾排放登记方式",
        "无源头减量措施",
        "容器垃圾不纯净",
        "无废弃油脂合同或不合格",
        "合计",
    ]
    metrics = [
        "无适量点餐、光盘行动等宣传内容",
        "集中用餐区未成组设置厨余和其他垃圾容器数",
        "容器无标识或标识不合格数",
        "无宣传氛围",
        "无厨余垃圾收运合同或不合格",
        "无其他垃圾收运合同或不合格",
        "无厨余垃圾排放登记方式",
        "无非居民其他垃圾排放登记方式",
        "无源头减量措施",
        "容器垃圾不纯净",
        "无废弃油脂合同或不合格",
    ]
    metric_values = [summary.values("餐饮单位检查情况", metric, required=False) for metric in metrics]
    total_values = summary.values("餐饮单位检查情况", "餐饮单位检查问题数")
    rows = [["总计", sum(counts.values()), *[fmt_int(sum(values.values())) for values in metric_values], fmt_int(sum(total_values.values()))]]
    for street, _value in ranked(total_values, reverse=True, count=len(STREETS), include_zero=True):
        rows.append([street, counts.get(street, 0), *[fmt_int(values[street]) for values in metric_values], fmt_int(total_values[street])])
    return headers, rows


def update_doc_tables(doc: Document, report_kind: str, summary: SummaryData, ledger_counts: dict[str, dict[str, int]]) -> None:
    if report_kind == "social" and len(doc.tables) >= 2:
        headers, rows = social_table_rows(summary, ledger_counts.get("社会单位", {}))
        fill_table(doc.tables[0], headers, rows)
        headers, rows = restaurant_table_rows(summary, ledger_counts.get("餐饮单位", {}))
        fill_table(doc.tables[1], headers, rows)


def make_replacements(
    summary: SummaryData,
    ledger_counts: dict[str, dict[str, int]],
    ledger_issue_totals: dict[str, list[tuple[str, float]]],
    start_date: datetime,
    end_date: datetime,
) -> dict[str, str]:
    period = format_period(start_date, end_date)
    month = format_month(end_date)
    construction = summary.values("分类设施建设达标情况", "合计")
    management = summary.values("分类设施管理达标情况", "合计")
    resident_bad = summary.values("居民自主投放情况", "居民自主投放不准确")
    total_by_street = {street: construction[street] + management[street] + resident_bad[street] for street in STREETS}
    total_issues = sum(total_by_street.values())
    top_streets = ranked(total_by_street, reverse=True, count=3)
    low_streets = ranked(total_by_street, reverse=False, count=3)
    main_issues = summary.top_issue_metrics(
        [
            ("居民自主投放情况", "居民自主投放不准确"),
            ("分类设施管理达标情况", "桶站周边不洁"),
            ("分类设施管理达标情况", "桶站地面脏污"),
            ("分类设施管理达标情况", "容器满冒"),
            ("分类设施管理达标情况", "垃圾积存"),
            ("分类设施管理达标情况", "桶内分类不纯净"),
            ("分类设施建设达标情况", "高峰时段未开盖"),
            ("分类设施建设达标情况", "合计"),
        ],
        count=6,
    )
    accuracy = summary.values("居民自主投放情况", "投放准确率")
    dirty = {
        street: summary.values("分类设施管理达标情况", "桶站周边不洁")[street]
        + summary.values("分类设施管理达标情况", "桶站地面脏污")[street]
        + summary.values("分类设施管理达标情况", "容器满冒")[street]
        + summary.values("分类设施管理达标情况", "垃圾积存")[street]
        for street in STREETS
    }
    peak = summary.values("分类设施建设达标情况", "高峰时段未开盖")
    social = summary.values("社会单位检查情况", "社会单位检查问题数")
    restaurant = summary.values("餐饮单位检查情况", "餐饮单位检查问题数")
    social_total = sum(social.values())
    restaurant_total = sum(restaurant.values())
    social_top = ranked(social, reverse=True, count=2, include_zero=False)
    restaurant_top = ranked(restaurant, reverse=True, count=2, include_zero=False)
    datasets = build_chart_datasets(summary, ledger_issue_totals)
    social_type_text = "、".join(name for name, _ in datasets["SOCIAL_ISSUE_TYPES"][:4]) or "暂无"
    restaurant_type_text = "、".join(name for name, _ in datasets["RESTAURANT_ISSUE_TYPES"][:6]) or "暂无"
    replacements = {
        "{{START_DATE}}": f"{start_date.year}年{start_date.month}月{start_date.day}日",
        "{{END_DATE}}": f"{end_date.year}年{end_date.month}月{end_date.day}日",
        "{{PERIOD}}": period,
        "{{MONTH}}": month,
        "{{TOTAL_ISSUES}}": fmt_int(total_issues),
        "{{TOP_STREETS_TEXT}}": with_ratio(top_streets, total_issues),
        "{{LOW_STREETS_TEXT}}": with_ratio(low_streets, total_issues),
        "{{MAIN_ISSUE_TYPES_TEXT}}": "、".join(f"{name}（{fmt_int(value)}处，占比{fmt_pct(value / sum(v for _, v in main_issues) if main_issues else 0)}）" for name, value in main_issues),
        "{{LOW_ACCURACY_TEXT}}": "、".join(f"{name}{fmt_pct(value)}" for name, value in ranked(accuracy, reverse=False, count=3)),
        "{{HIGH_ACCURACY_TEXT}}": "、".join(f"{name}{fmt_pct(value)}" for name, value in ranked(accuracy, reverse=True, count=3)),
        "{{DIRTY_TOP_TEXT}}": "、".join(f"{name}{fmt_int(value)}处" for name, value in ranked(dirty, reverse=True, count=3, include_zero=False)),
        "{{PEAK_TOP_TEXT}}": "、".join(f"{name}{fmt_int(value)}处" for name, value in ranked(peak, reverse=True, count=3, include_zero=False)),
        "{{SOCIAL_UNIT_COUNT}}": fmt_int(sum(ledger_counts.get("社会单位", {}).values())),
        "{{RESTAURANT_UNIT_COUNT}}": fmt_int(sum(ledger_counts.get("餐饮单位", {}).values())),
        "{{SOCIAL_ISSUES}}": fmt_int(social_total),
        "{{RESTAURANT_ISSUES}}": fmt_int(restaurant_total),
        "{{SOCIAL_RESTAURANT_ISSUES}}": fmt_int(social_total + restaurant_total),
        "{{SOCIAL_TOP_TEXT}}": with_ratio(social_top, social_total) if social_top else "暂无",
        "{{RESTAURANT_TOP_TEXT}}": with_ratio(restaurant_top, restaurant_total) if restaurant_top else "暂无",
        "{{SOCIAL_MAIN_TYPES_TEXT}}": social_type_text,
        "{{RESTAURANT_MAIN_TYPES_TEXT}}": restaurant_type_text,
    }
    # Sample-document fallbacks.
    replacements.update(
        {
            "2026年2月17日至2026年3月19日": period,
            "2026年3月份": month,
            "共计发现问题955处次": f"共计发现问题{fmt_int(total_issues)}处次",
            "发现问题955处次": f"发现问题{fmt_int(total_issues)}处次",
            "共计发现问题109处次": f"共计发现问题{fmt_int(social_total + restaurant_total)}处次",
            "共计发现问题45处次": f"共计发现问题{fmt_int(social_total)}处次",
            "共计发现问题64处次": f"共计发现问题{fmt_int(restaurant_total)}处次",
            "91个社会单位": f"{fmt_int(sum(ledger_counts.get('社会单位', {}).values()))}个社会单位",
            "155个餐饮单位": f"{fmt_int(sum(ledger_counts.get('餐饮单位', {}).values()))}个餐饮单位",
        }
    )
    return replacements


def set_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    for run in paragraph.runs:
        run.text = ""
    paragraph.runs[0].text = text


def update_narrative(
    doc: Document,
    report_kind: str,
    summary: SummaryData,
    ledger_counts: dict[str, dict[str, int]],
    ledger_issue_totals: dict[str, list[tuple[str, float]]],
    start_date: datetime,
    end_date: datetime,
) -> None:
    period = format_period(start_date, end_date)
    construction = summary.values("分类设施建设达标情况", "合计")
    management = summary.values("分类设施管理达标情况", "合计")
    resident_bad = summary.values("居民自主投放情况", "居民自主投放不准确")
    total_by_street = {street: construction[street] + management[street] + resident_bad[street] for street in STREETS}
    total_issues = sum(total_by_street.values())
    social = summary.values("社会单位检查情况", "社会单位检查问题数")
    restaurant = summary.values("餐饮单位检查情况", "餐饮单位检查问题数")
    social_total = sum(social.values())
    restaurant_total = sum(restaurant.values())
    social_counts = ledger_counts.get("社会单位", {})
    restaurant_counts = ledger_counts.get("餐饮单位", {})
    datasets = build_chart_datasets(summary, ledger_issue_totals)

    paragraph_map: list[tuple[str, str]] = []
    if report_kind == "social":
        social_top = ranked(social, reverse=True, count=2, include_zero=False)
        restaurant_top = ranked(restaurant, reverse=True, count=2, include_zero=False)
        social_leader = social_top[0] if social_top else ("暂无", 0)
        social_runner = social_top[1] if len(social_top) > 1 else ("暂无", 0)
        restaurant_leader = restaurant_top[0] if restaurant_top else ("暂无", 0)
        restaurant_runner = restaurant_top[1] if len(restaurant_top) > 1 else ("暂无", 0)
        paragraph_map = [
            (
                "对本区15个街道",
                f"{period}对本区15个街道{sum(social_counts.values())}个社会单位和{sum(restaurant_counts.values())}个餐饮单位的垃圾分类的设施建设和管理情况进行了抽查，共计发现问题{fmt_int(social_total + restaurant_total)}处次，问题以台账的形式发给各街道办事处，各街道办事处整改后以书面形式回复整改结果，第三方检查组将对整改结果进行审核。",
            ),
            ("本周期共计检查", f"本周期共计检查{sum(social_counts.values())}个社会单位，共计发现问题{fmt_int(social_total)}处次。"),
            (
                "本月检查社会单位中",
                f"本月检查社会单位中，问题较多的街道是{social_leader[0]}，检查出{fmt_int(social_leader[1])}处问题；其次是{social_runner[0]}，检查出{fmt_int(social_runner[1])}处问题。",
            ),
            (
                "从问题类型来看，社会单位",
                f"从问题类型来看，社会单位的问题主要出现在{'、'.join(name for name, _ in datasets['SOCIAL_ISSUE_TYPES'][:4]) or '暂无'}。",
            ),
            ("本周期共计检查", f"本周期共计检查{sum(restaurant_counts.values())}个餐饮单位，共计发现问题{fmt_int(restaurant_total)}处次。"),
            (
                "本月检查餐饮单位中",
                f"本月检查餐饮单位中，问题较多的是{restaurant_leader[0]}，检查出{fmt_int(restaurant_leader[1])}处问题；其次是{restaurant_runner[0]}，检查出{fmt_int(restaurant_runner[1])}处问题。",
            ),
            (
                "从问题类型来看，餐饮单位",
                f"从问题类型来看，餐饮单位的问题主要出现在{'、'.join(name for name, _ in datasets['RESTAURANT_ISSUE_TYPES'][:6]) or '暂无'}。",
            ),
        ]
        used_cycle_count = 0
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if text.startswith("本周期共计检查"):
                replacement = paragraph_map[1][1] if used_cycle_count == 0 else paragraph_map[4][1]
                used_cycle_count += 1
                set_paragraph_text(paragraph, replacement)
                continue
            for needle, replacement in paragraph_map:
                if needle != "本周期共计检查" and needle in text:
                    set_paragraph_text(paragraph, replacement)
                    break
        return

    social = summary.values("社会单位检查情况", "社会单位检查问题数")
    restaurant = summary.values("餐饮单位检查情况", "餐饮单位检查问题数")
    social_total = sum(social.values())
    restaurant_total = sum(restaurant.values())
    social_top = ranked(social, reverse=True, count=2, include_zero=False)
    restaurant_top = ranked(restaurant, reverse=True, count=2, include_zero=False)
    accuracy = summary.values("居民自主投放情况", "投放准确率")
    purity = summary.values("纯净率", "桶内分类纯净率")
    dirty = {street: value for street, value in datasets["DIRTY_OVERFLOW"]}
    peak = summary.values("分类设施建设达标情况", "高峰时段未开盖")
    paragraph_map = [
        (
            "对本区15个街道",
            f"{period}对本区15个街道的居住小区、平房区垃圾分类设施建设、日常管理、居民自主投放等情况进行了检查，共计发现问题{fmt_int(total_issues)}处次，问题以日报的形式发给各街道办事处，各街道办事处整改后以书面形式回复整改结果。",
        ),
        (
            "本周期对15个街道进行了全覆盖检查",
            f"本周期对15个街道进行了全覆盖检查，发现问题{fmt_int(total_issues)}处次。问题较多的街道依次是{with_ratio(ranked(total_by_street, reverse=True, count=3), total_issues)}。问题较少的街道依次是{with_ratio(ranked(total_by_street, reverse=False, count=3), total_issues)}。",
        ),
        (
            "从本周期检查发现问题情况看",
            f"从本周期检查发现问题情况看，检查发现问题类型如下：{'、'.join(f'{name}（{fmt_int(value)}处）' for name, value in datasets['MAIN_ISSUE_TYPES'])}。",
        ),
        (
            "本周期居民自主分类投放准确率",
            f"本周期居民自主分类投放准确率较低的街道为{'、'.join(f'{name}，{fmt_pct(value)}' for name, value in ranked(accuracy, reverse=False, count=3))}；较高的街道为{'、'.join(f'{name}，{fmt_pct(value)}' for name, value in ranked(accuracy, reverse=True, count=3))}。",
        ),
        (
            "本周期桶内垃圾分类纯净率",
            f"本周期桶内垃圾分类纯净率较低的街道为{'、'.join(f'{name}，{fmt_pct(value)}' for name, value in ranked(purity, reverse=False, count=3))}；较高的街道为{'、'.join(f'{name}，{fmt_pct(value)}' for name, value in ranked(purity, reverse=True, count=3))}。",
        ),
        (
            "本周期桶站满冒环境脏乱数较多",
            f"本周期桶站满冒环境脏乱数较多的街道为{list_counts(ranked(dirty, reverse=True, count=3, include_zero=False))}；较少的街道为{list_counts(ranked(dirty, reverse=False, count=3))}。",
        ),
        (
            "本周期高峰时段未开盖数较多",
            f"本周期高峰时段未开盖数较多的街道为{list_counts(ranked(peak, reverse=True, count=3, include_zero=False))}；较少的街道为{list_counts(ranked(peak, reverse=False, count=3))}。",
        ),
        (
            "本周期桶站建设不规范问题较多",
            f"本周期桶站建设不规范问题较多的街道为{list_counts(ranked(construction, reverse=True, count=3, include_zero=False))}；问题较少的街道为{list_counts(ranked(construction, reverse=False, count=3))}。",
        ),
        (
            "对社会单位和餐饮单位开展检查工作",
            f"{period}，对社会单位和餐饮单位开展检查工作，检查发现社会单位问题{fmt_int(social_total)}个，问题较多的街道为{ratio_items(social_top, social_total)}；餐饮单位问题{fmt_int(restaurant_total)}个，问题较多的街道为{ratio_items(restaurant_top, restaurant_total)}。",
        ),
        (
            "社会单位存在问题较多的为",
            f"社会单位存在问题较多的为：{'、'.join(name for name, _ in datasets['SOCIAL_ISSUE_TYPES'][:6]) or '暂无'}；",
        ),
        (
            "餐饮单位存在问题较多的为",
            f"餐饮单位存在问题较多的为：{'、'.join(name for name, _ in datasets['RESTAURANT_ISSUE_TYPES'][:6]) or '暂无'}。",
        ),
    ]
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for needle, replacement in paragraph_map:
            if needle in text:
                set_paragraph_text(paragraph, replacement)
                break


def chart_rel_target(chart_name: str, members: dict[str, bytes]) -> str | None:
    rel_name = f"word/charts/_rels/{Path(chart_name).name}.rels"
    if rel_name not in members:
        return None
    root = ET.fromstring(members[rel_name])
    for rel in root:
        target = rel.attrib.get("Target", "")
        if "embeddings/" in target and target.endswith(".xlsx"):
            return "word/" + target.split("word/")[-1].replace("../", "")
    return None


def update_embedded_workbook(blob: bytes, data: list[tuple[str, float]]) -> bytes:
    wb = load_workbook(io.BytesIO(blob))
    ws = wb.active
    ws.delete_rows(1, ws.max_row)
    ws["A1"] = "项目"
    ws["B1"] = "数值"
    for row_idx, (label, value) in enumerate(data, start=2):
        ws.cell(row_idx, 1, label)
        ws.cell(row_idx, 2, float(value))
    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


def set_cache(parent, values: list[str], cache_tag: str) -> None:
    cache = parent.find(f"c:{cache_tag}", CHART_NS)
    if cache is None:
        cache = ET.SubElement(parent, f"{{{CHART_NS['c']}}}{cache_tag}")
    for pt_count in cache.findall("c:ptCount", CHART_NS):
        cache.remove(pt_count)
    for pt in cache.findall("c:pt", CHART_NS):
        cache.remove(pt)
    pt_count = ET.SubElement(cache, f"{{{CHART_NS['c']}}}ptCount")
    pt_count.set("val", str(len(values)))
    for idx, value in enumerate(values):
        pt = ET.SubElement(cache, f"{{{CHART_NS['c']}}}pt")
        pt.set("idx", str(idx))
        v = ET.SubElement(pt, f"{{{CHART_NS['c']}}}v")
        v.text = value


def update_chart_xml(blob: bytes, data: list[tuple[str, float]]) -> bytes:
    root = ET.fromstring(blob)
    cats = [label for label, _value in data]
    vals = [str(float(value)) for _label, value in data]
    end_row = len(data) + 1
    cat_ref = f"Sheet1!$A$2:$A${end_row}"
    val_ref = f"Sheet1!$B$2:$B${end_row}"
    for ser in root.findall(".//c:ser", CHART_NS):
        cat = ser.find("c:cat", CHART_NS)
        val = ser.find("c:val", CHART_NS)
        if cat is not None:
            str_ref = cat.find("c:strRef", CHART_NS)
            if str_ref is None:
                str_ref = ET.SubElement(cat, f"{{{CHART_NS['c']}}}strRef")
            f = str_ref.find("c:f", CHART_NS)
            if f is None:
                f = ET.SubElement(str_ref, f"{{{CHART_NS['c']}}}f")
            f.text = cat_ref
            set_cache(str_ref, cats, "strCache")
        if val is not None:
            num_ref = val.find("c:numRef", CHART_NS)
            if num_ref is None:
                num_ref = ET.SubElement(val, f"{{{CHART_NS['c']}}}numRef")
            f = num_ref.find("c:f", CHART_NS)
            if f is None:
                f = ET.SubElement(num_ref, f"{{{CHART_NS['c']}}}f")
            f.text = val_ref
            set_cache(num_ref, vals, "numCache")
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def update_charts(docx_path: Path, chart_keys: list[tuple[str, str]], datasets: dict[str, list[tuple[str, float]]]) -> None:
    tmp = docx_path.with_suffix(".charts.tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as zin:
        members = {name: zin.read(name) for name in zin.namelist()}
    chart_names = sorted(
        [name for name in members if re.fullmatch(r"word/charts/chart\d+\.xml", name)],
        key=lambda name: int(re.search(r"chart(\d+)\.xml", name).group(1)),
    )
    for chart_name, (dataset_key, _label) in zip(chart_names, chart_keys):
        data = datasets.get(dataset_key, [])
        if not data:
            continue
        members[chart_name] = update_chart_xml(members[chart_name], data)
        target = chart_rel_target(chart_name, members)
        if target and target in members:
            members[target] = update_embedded_workbook(members[target], data)
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, blob in members.items():
            zout.writestr(name, blob)
    replace_file_with_retry(tmp, docx_path)


def replace_file_with_retry(source: Path, target: Path, attempts: int = 5) -> None:
    last_error: PermissionError | None = None
    for _ in range(attempts):
        try:
            source.replace(target)
            return
        except PermissionError as exc:
            last_error = exc
            time_module.sleep(0.4)
    raise PermissionError(
        f"无法覆盖 {target}。请关闭已经打开的同名报告或预览窗口后重试。"
    ) from last_error


def available_output_path(path: Path) -> Path:
    if not path.exists():
        return path
    stem = path.stem
    suffix = path.suffix
    for index in range(1, 100):
        candidate = path.with_name(f"{stem}_{index}{suffix}")
        if not candidate.exists():
            return candidate
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    return path.with_name(f"{stem}_{timestamp}{suffix}")


def build_report(
    template_docx: Path,
    output_docx: Path,
    report_kind: str,
    summary: SummaryData,
    ledger_counts: dict[str, dict[str, int]],
    ledger_issue_totals: dict[str, list[tuple[str, float]]],
    start_date: datetime,
    end_date: datetime,
) -> None:
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template_docx, output_docx)
    doc = Document(output_docx)
    replace_text(doc, make_replacements(summary, ledger_counts, ledger_issue_totals, start_date, end_date))
    update_narrative(doc, report_kind, summary, ledger_counts, ledger_issue_totals, start_date, end_date)
    update_doc_tables(doc, report_kind, summary, ledger_counts)
    doc.save(output_docx)
    datasets = build_chart_datasets(summary, ledger_issue_totals)
    chart_keys = RESIDENTIAL_CHART_METRICS if report_kind == "residential" else SOCIAL_CHART_METRICS
    update_charts(output_docx, chart_keys, datasets)


def generate_reports(
    summary_xlsx: str | Path,
    source_ledger_xlsx: str | Path,
    residential_template_docx: str | Path,
    social_template_docx: str | Path,
    start_date: datetime,
    end_date: datetime,
    output_dir: str | Path,
) -> dict:
    summary = SummaryData.load(Path(summary_xlsx))
    ledger_counts = load_ledger_counts(Path(source_ledger_xlsx), start_date, end_date)
    ledger_issue_totals = load_ledger_issue_totals(Path(source_ledger_xlsx), start_date, end_date)
    output_dir = Path(output_dir)
    period = f"{start_date.year}.{start_date.month}.{start_date.day}-{end_date.year}.{end_date.month}.{end_date.day}"
    residential_output = available_output_path(output_dir / f"西城区{period}垃圾分类检查自查报告.docx")
    social_output = available_output_path(output_dir / f"西城区{period}社会单位、餐饮单位检查报告.docx")
    build_report(Path(residential_template_docx), residential_output, "residential", summary, ledger_counts, ledger_issue_totals, start_date, end_date)
    build_report(Path(social_template_docx), social_output, "social", summary, ledger_counts, ledger_issue_totals, start_date, end_date)
    return {
        "output_files": [residential_output, social_output],
        "missing": sorted(set(summary.missing)),
        "ledger_counts": ledger_counts,
        "ledger_issue_totals": ledger_issue_totals,
    }


def make_template(source_docx: Path, output_docx: Path, kind: str) -> None:
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source_docx, output_docx)
    doc = Document(output_docx)
    replacements = {
        "2026年2月17日至2026年3月19日": "{{PERIOD}}",
        "2026年3月份": "{{MONTH}}",
    }
    if kind == "residential":
        replacements.update(
            {
                "共计发现问题955处次": "共计发现问题{{TOTAL_ISSUES}}处次",
                "发现问题955处次": "发现问题{{TOTAL_ISSUES}}处次",
            }
        )
    else:
        replacements.update(
            {
                "检查91个社会单位": "检查{{SOCIAL_UNIT_COUNT}}个社会单位",
                "检查155个餐饮单位": "检查{{RESTAURANT_UNIT_COUNT}}个餐饮单位",
                "共计发现问题109处次": "共计发现问题{{SOCIAL_RESTAURANT_ISSUES}}处次",
                "共计发现问题45处次": "共计发现问题{{SOCIAL_ISSUES}}处次",
                "共计发现问题64处次": "共计发现问题{{RESTAURANT_ISSUES}}处次",
            }
        )
    replace_text(doc, replacements)
    doc.save(output_docx)


def create_default_templates(residential_sample: Path, social_sample: Path, output_dir: Path = DEFAULT_TEMPLATE_DIR) -> dict:
    residential_template = output_dir / "垃圾分类检查自查报告模板.docx"
    social_template = output_dir / "社会单位、餐饮单位检查报告模板.docx"
    make_template(residential_sample, residential_template, "residential")
    make_template(social_sample, social_template, "social")
    return {"residential_template": residential_template, "social_template": social_template}


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate DOCX monthly garbage-classification reports.")
    parser.add_argument("--summary-xlsx", required=True, type=Path)
    parser.add_argument("--source-ledger-xlsx", required=True, type=Path)
    parser.add_argument("--residential-template", required=True, type=Path)
    parser.add_argument("--social-template", required=True, type=Path)
    parser.add_argument("--start-date", required=True)
    parser.add_argument("--end-date", required=True)
    parser.add_argument("--output-dir", required=True, type=Path)
    args = parser.parse_args()
    result = generate_reports(
        args.summary_xlsx,
        args.source_ledger_xlsx,
        args.residential_template,
        args.social_template,
        parse_date(args.start_date),
        parse_date(args.end_date, end_of_day=True),
        args.output_dir,
    )
    for path in result["output_files"]:
        print(path)
    if result["missing"]:
        print("缺失字段：")
        for item in result["missing"]:
            print(f"- {item}")


if __name__ == "__main__":
    main()
